import os, json
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import brand_doc as B
import config as C
import lab_parser as LP
import paths as P

SP = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = P.OUT
man = json.load(open(P.MANIFEST))
LABS = {L['num']: L for L in LP.load_all(P.LABS)}

TOTAL = man['total_slides']
lab_slide = {int(k): v['slide'] for k, v in man['labs'].items()}
anchors = sorted(list(lab_slide.values()) + [d['section_slide'] for d in man['domains']])

def lab_range(ln):
    s = lab_slide[ln]
    nxt = [a for a in anchors if a > s]
    e = (min(nxt) - 1) if nxt else TOTAL - 8
    return s, e

def rng(a, b): return str(a) if a == b else f"{a}–{b}"
def short(ln):
    t = LABS[ln]['title']; return t.split('—',1)[-1].strip()
def span(labnums):
    s, _ = lab_range(labnums[0]); _, e = lab_range(labnums[-1]); return rng(s, e)

doc = B.new_doc()

def center_img(path, w):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(path, width=Inches(w))
def center_text(t, s, b=False, col=None, sp=4):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(sp)
    r = p.add_run(t); r.font.size = Pt(s); r.font.bold = b; r.font.name='Arial'
    if col: r.font.color.rgb = col

# ---- Cover ----
doc.add_paragraph().paragraph_format.space_after = Pt(2)
center_img(os.path.join(B.ASSETS, "tertiary_logo.png"), 1.5)
center_text("Tertiary Infotech Academy Pte Ltd", 13, True, B.INK, 1)
center_text("UEN: 201200696W", 9.5, False, B.GREY, 10)
center_text("LESSON PLAN", 26, True, B.BLUE, 2)
center_text("For", 11, False, B.GREY, 10)
center_img(os.path.join(B.ASSETS, "cysa_logo.png"), 1.3)
doc.add_paragraph().paragraph_format.space_after = Pt(2)
center_text(C.COURSE_TITLE, 20, True, B.INK, 4)
center_text(f"TGS Ref No: {C.COURSE_CODE}", 11, False, B.INK, 10)
center_text(f"Version {C.VERSION_DOC}", 12, True, B.BLUE, 4)
doc.add_page_break()

# ---- Version control ----
B.heading(doc, "Document Version Control Record", 1)
B.blue_table(doc, ["Version", "Effective Date", "Summary of Changes", "Author"], [
    ["1.0", "6 July 2026",
     "First edition — 5-day lesson plan aligned to the 30 hands-on CySA+ labs, with slide-number references to the Learner Guide Slides deck.",
     "Tertiary Infotech Academy Pte Ltd"],
], widths=[Inches(0.8), Inches(1.2), Inches(3.6), Inches(1.4)])
doc.add_page_break()

# ---- Course details ----
B.heading(doc, "Course Details", 1)
B.blue_table(doc, ["Field", "Detail"], [
    ["Course Title", C.COURSE_TITLE],
    ["Course Code", C.COURSE_CODE],
    ["WSQ TSC", f"{C.TSC_TITLE} ({C.TSC_CODE})"],
    ["Duration", "5 Days (40 hours)"],
    ["Daily Schedule", "9:00 AM – 6:00 PM (8 training hours/day, excluding lunch)"],
    ["Lunch Break", "12:30 PM – 1:15 PM (45 minutes)"],
    ["Breaks", "Short tea breaks are scheduled within each day's training hours"],
    ["Delivery Mode", "Instructor-led presentation, live demonstration and hands-on labs (Killercoda / free tools)"],
    ["Assessment", "Written Assessment (SAQ) 1 hr + Case Study (CS) 1 hr — Open Book (Day 5)"],
    ["Prerequisites", "Basic networking and Linux command-line familiarity; interest in cybersecurity operations"],
    ["Slide Deck", f"{man['ppt_file']} ({TOTAL} slides)"],
], widths=[Inches(1.8), Inches(4.8)], label_col0=True)
B.para(doc, "Note: the 'Slides' column in each daily schedule references the exact slide numbers in the Learner Guide Slides deck for that activity.",
       italic=True, color=B.GREY, size=9.5)

# ---- Overview + LO ----
B.heading(doc, "Course Overview", 1)
B.para(doc, "This 5-day hands-on course develops practical cybersecurity analyst skills aligned to the CompTIA CySA+ "
            "CS0-003 exam. Across four domains — Security Operations, Vulnerability Management, Incident Response, and "
            "Reporting & Communication — learners complete 30 step-by-step labs using the free Killercoda playground and "
            "free industry tools, culminating in a written and case-study assessment.")
B.heading(doc, "Learning Outcomes", 1)
B.para(doc, "By the end of this course, participants will be able to:", bold=True, space=2)
B.bullets(doc, C.LEARNING_OUTCOMES)
doc.add_page_break()

# ---- Daily schedule tables ----
HDR = ["Time", "Topic / Activity", "Duration", "Slides"]
WID = [Inches(1.15), Inches(3.5), Inches(0.85), Inches(0.9)]

def sched_table(rows):
    t = doc.add_table(rows=1, cols=4); t.style = 'Table Grid'
    for c, h in zip(t.rows[0].cells, HDR):
        B.shade(c._tc.get_or_add_tcPr(), '1F6FEB'); c.paragraphs[0].text = ''
        r = c.paragraphs[0].add_run(h); r.font.bold = True; r.font.color.rgb = B.WHITE; r.font.size = Pt(9.5); r.font.name='Arial'
    for row in rows:
        kind = row[0]; vals = row[1:]
        cells = t.add_row().cells
        fill = {'topic': 'EAF1FD', 'break': 'FDF3E7'}.get(kind)
        for i, (cell, v) in enumerate(zip(cells, vals)):
            cell.paragraphs[0].text = ''
            r = cell.paragraphs[0].add_run(str(v)); r.font.size = Pt(9); r.font.name='Arial'
            if kind == 'topic': r.font.bold = True; r.font.color.rgb = B.BLUE
            elif kind == 'break': r.font.italic = True; r.font.color.rgb = B.GREY
            else: r.font.color.rgb = B.INK
            if fill: B.shade(cell._tc.get_or_add_tcPr(), fill)
    for row in t.rows:
        for i, wd in enumerate(WID): row.cells[i].width = wd
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def labs_rows(labnums, blocks):
    """distribute labnums across content 'blocks' [(time,dur)] -> rows list."""
    n = len(blocks); groups = [[] for _ in range(n)]
    for i, ln in enumerate(labnums): groups[i * n // len(labnums)].append(ln)
    rows = []
    for (time, dur), grp in zip(blocks, groups):
        if not grp: continue
        act = "; ".join(f"Lab {x}: {short(x)}" for x in grp)
        rows.append(('normal', time, act, dur, span(grp)))
    return rows

def day(dayno, subtitle, labnums, admin=False):
    B.heading(doc, f"Day {dayno}", 1)
    B.para(doc, subtitle, italic=True, color=B.GREY, space=6)
    dom = next(d for d in man['domains'] if d['labs'] == labnums) if labnums else None
    # Morning
    B.para(doc, "Morning Session  ·  9:00 AM – 12:30 PM", bold=True, color=B.BLUE, space=3)
    rows = []
    if admin:
        rows += [('normal', "9:00 – 9:20", "Digital Attendance (AM); Trainer & Learner Introduction; Ground Rules", "20 min", "1–5"),
                 ('normal', "9:20 – 9:55", "Lesson Plan; Learning Outcomes; Skills Framework; Course Outline", "35 min", "6–19"),
                 ('normal', "9:55 – 10:15", "Final Assessment Briefing; Criteria for Funding; LMS/TMS", "20 min", "20–22"),
                 ('break', "10:15 – 10:30", "Tea Break", "15 min", "—")]
        am_blocks = [("10:30 – 12:30", "120 min")]
    else:
        rows += [('normal', "9:00 – 9:10", "Digital Attendance (AM); Recap of previous day", "10 min", "—")]
        am_blocks = [("9:10 – 10:30", "80 min"), ("10:45 – 12:30", "105 min")]
    if dom:
        rows.append(('topic', "", f"{dom['topic']}: {dom['name']} — Topic Overview", "—", rng(dom['section_slide'], dom['section_slide']+1)))
    half = (len(labnums) + 1) // 2
    am_rows = labs_rows(labnums[:half], am_blocks)
    if not admin and len(am_blocks) > 1:
        # insert tea break between the two am blocks
        am_rows = [am_rows[0]] + [('break', "10:30 – 10:45", "Tea Break", "15 min", "—")] + am_rows[1:]
    rows += am_rows
    rows.append(('break', "12:30 – 1:15", "Lunch Break", "45 min", "—"))
    sched_table(rows)
    # Afternoon
    B.para(doc, "Afternoon Session  ·  1:15 PM – 6:00 PM", bold=True, color=B.BLUE, space=3)
    rows2 = [('normal', "1:15 – 1:20", "Digital Attendance (PM)", "5 min", "—")]
    pm_blocks = [("1:20 – 3:30", "130 min"), ("3:45 – 6:00", "135 min")]
    pm_rows = labs_rows(labnums[half:], pm_blocks)
    if pm_rows:
        rows2.append(pm_rows[0]); rows2.append(('break', "3:30 – 3:45", "Tea Break", "15 min", "—"))
        rows2 += pm_rows[1:]
    rows2.append(('normal', "6:00", f"End of Day {dayno}", "—", "—"))
    sched_table(rows2)
    doc.add_page_break()

DOMS = [(1, "Topic 1: Security Operations", list(range(1,11)), True),
        (2, "Topic 2: Vulnerability Management", list(range(11,20)), False),
        (3, "Topic 3: Incident Response and Management", list(range(20,26)), False),
        (4, "Topic 4: Reporting and Communication", list(range(26,31)), False)]
for dn, sub, labs, adm in DOMS:
    day(dn, sub, labs, admin=adm)

# Day 5
B.heading(doc, "Day 5", 1)
B.para(doc, "Revision and Final Assessment", italic=True, color=B.GREY, space=6)
B.para(doc, "Morning Session  ·  9:00 AM – 12:30 PM", bold=True, color=B.BLUE, space=3)
back = TOTAL - 8 + 1
sched_table([
    ('normal', "9:00 – 9:10", "Digital Attendance (AM)", "10 min", "—"),
    ('normal', "9:10 – 10:30", "Revision: Topic 1 Security Operations & Topic 2 Vulnerability Management", "80 min",
        f"{span(list(range(1,11)))}, {span(list(range(11,20)))}"),
    ('break', "10:30 – 10:45", "Tea Break", "15 min", "—"),
    ('normal', "10:45 – 12:30", "Revision: Topic 3 Incident Response & Topic 4 Reporting; Practice Exams", "105 min",
        f"{span(list(range(20,26)))}, {span(list(range(26,31)))}, {rng(back,back)}"),
    ('break', "12:30 – 1:15", "Lunch Break", "45 min", "—"),
])
B.para(doc, "Afternoon Session  ·  1:15 PM – 6:00 PM", bold=True, color=B.BLUE, space=3)
sched_table([
    ('normal', "1:15 – 1:20", "Digital Attendance (PM)", "5 min", "—"),
    ('normal', "1:20 – 3:30", "Q&A; Summary; exam techniques; Course Feedback & TRAQOM Survey", "130 min", rng(back+1, back+3)),
    ('break', "3:30 – 3:45", "Tea Break", "15 min", "—"),
    ('normal', "3:45 – 4:00", "Assessment Briefing; Digital Attendance (Assessment)", "15 min", rng(back+4, back+4)),
    ('normal', "4:00 – 6:00", "Final Assessment: Written (SAQ) 1 hr + Case Study (CS) 1 hr", "120 min", "—"),
    ('normal', "6:00", "End of Class; Recommended Courses; Support", "—", rng(back+5, TOTAL)),
])

# ---- Slide reference summary ----
doc.add_page_break()
B.heading(doc, "Slide Reference Summary", 1)
rows = [["Front Admin", "Title, Attendance, Ground Rules, Lesson Plan, Skills Framework, Learning Outcomes, Course Outline, Assessment, Funding, LMS", "1–22"]]
for d in man['domains']:
    rows.append([f"{d['topic']}: {d['name']}", f"Section + Overview + Labs {d['labs'][0]}–{d['labs'][-1]}", span(d['labs'])])
for ln in range(1,31):
    s,e = lab_range(ln); rows.append([f"Lab {ln}", short(ln), rng(s,e)])
rows.append(["Back Admin", "Practice Exams, Summary/Q&A, TRAQOM, Attendance, Final Assessment, Recommended Courses, Support, Thank You", rng(back, TOTAL)])
B.blue_table(doc, ["Section", "Content", "Slides"], rows, widths=[Inches(1.9), Inches(3.9), Inches(0.9)])

B.add_footer(doc)
lp_name = f"LP-{C.COURSE_TITLE}.docx"
doc.save(os.path.join(OUT_DIR, lp_name))
print("Saved LP:", lp_name)
print("lab ranges:", {ln: lab_range(ln) for ln in [1,10,11,19,20,25,26,30]})
