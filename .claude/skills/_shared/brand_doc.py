"""Shared branded DOCX helpers (Tertiary Infotech style)."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SP = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(SP, "assets")

BLUE = RGBColor(0x1F, 0x6F, 0xEB)
GREEN = RGBColor(0x0F, 0xB9, 0x80)
INK = RGBColor(0x16, 0x1B, 0x26)
GREY = RGBColor(0x5B, 0x63, 0x72)
RED = RGBColor(0xC0, 0x2B, 0x2B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def new_doc():
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Arial'; st.font.size = Pt(10.5); st.font.color.rgb = INK
    for lvl, sz in [(1, 16), (2, 13.5), (3, 11.5)]:
        s = doc.styles[f'Heading {lvl}']
        s.font.name = 'Arial'; s.font.size = Pt(sz); s.font.bold = True; s.font.color.rgb = BLUE
    # widen margins a touch
    for sec in doc.sections:
        sec.left_margin = Inches(0.9); sec.right_margin = Inches(0.9)
        sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.8)
    return doc

def _field(run, instr):
    r = run._r
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    r.append(f1); r.append(it); r.append(f2)

def add_footer(doc, copy_line="© 2026 Tertiary Infotech Academy Pte Ltd. All rights reserved.  ·  www.tertiarycourses.com.sg"):
    for sec in doc.sections:
        f = sec.footer
        f.is_linked_to_previous = False
        p = f.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Page "); r.font.size = Pt(8.5); r.font.color.rgb = GREY; r.font.name='Arial'
        r2 = p.add_run(); r2.font.size = Pt(8.5); r2.font.color.rgb = GREY; _field(r2, "PAGE")
        r3 = p.add_run(" of "); r3.font.size = Pt(8.5); r3.font.color.rgb = GREY; r3.font.name='Arial'
        r4 = p.add_run(); r4.font.size = Pt(8.5); r4.font.color.rgb = GREY; _field(r4, "NUMPAGES")
        p2 = f.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = p2.add_run(copy_line); rr.font.size = Pt(7.5); rr.font.color.rgb = GREY; rr.font.name='Arial'

def shade(el, fill):
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), fill); el.append(sh)

def para(doc, text, size=10.5, bold=False, color=None, italic=False, space=6, align=None, mono=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space)
    if align: p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    r.font.name = 'Consolas' if mono else 'Arial'
    if color: r.font.color.rgb = color
    return p

def bullets(doc, items, style='List Bullet'):
    for it in items:
        p = doc.add_paragraph(style=style)
        run = p.add_run(it); run.font.size = Pt(10.5); run.font.name='Arial'; run.font.color.rgb = INK

def code_block(doc, text):
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
        shade(p._p.get_or_add_pPr(), 'F2F4F7')
        r = p.add_run(line if line else ' ')
        r.font.name = 'Consolas'; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x0B, 0x30, 0x60)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def note_box(doc, text, label="Note:"):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    shade(p._p.get_or_add_pPr(), 'FDF3E7')
    p.paragraph_format.left_indent = Inches(0.1); p.paragraph_format.right_indent = Inches(0.1)
    r1 = p.add_run(label + " "); r1.font.bold = True; r1.font.color.rgb = RGBColor(0xC8,0x7A,0x1E); r1.font.size = Pt(10.5); r1.font.name='Arial'
    r2 = p.add_run(text); r2.font.size = Pt(10.5); r2.font.color.rgb = INK; r2.font.name='Arial'

def blue_table(doc, headers, rows, widths=None, label_col0=False):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, h in zip(t.rows[0].cells, headers):
        shade(c._tc.get_or_add_tcPr(), '1F6FEB'); c.paragraphs[0].text = ''
        r = c.paragraphs[0].add_run(h); r.font.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(10); r.font.name='Arial'
    for row in rows:
        cells = t.add_row().cells
        for i, (cell, v) in enumerate(zip(cells, row)):
            cell.paragraphs[0].text = ''
            r = cell.paragraphs[0].add_run(str(v)); r.font.size = Pt(9.5); r.font.name='Arial'; r.font.color.rgb = INK
            if label_col0 and i == 0:
                shade(cell._tc.get_or_add_tcPr(), 'EEF3FC'); r.font.bold = True
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

def heading(doc, text, level, color=None):
    h = doc.add_heading(level=level)
    r = h.add_run(text); r.font.name = 'Arial'
    if color: r.font.color.rgb = color
    return h

def image_centered(doc, path, max_width_in=5.6):
    if not os.path.exists(path):
        return
    from PIL import Image
    iw, ih = Image.open(path).size
    w = min(max_width_in, 6.5)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(path, width=Inches(w))
