import os
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import brand_doc as B
import config as C
import lab_parser as LP
import paths as P

SP = os.path.dirname(os.path.abspath(__file__))
IMG = P.IMAGES
LABS = {L['num']: L for L in LP.load_all(P.LABS)}
OUT_DIR = P.OUT

doc = B.new_doc()

def center_img(path, width_in):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(path, width=Inches(width_in))

def center_text(text, size, bold=False, color=None, space=4):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text); r.font.size = Pt(size); r.font.bold = bold; r.font.name='Arial'
    if color: r.font.color.rgb = color

# ---------- Cover ----------
doc.add_paragraph().paragraph_format.space_after = Pt(2)
center_img(os.path.join(B.ASSETS, "tertiary_logo.png"), 1.5)
center_text("Tertiary Infotech Academy Pte Ltd", 13, True, B.INK, 1)
center_text("UEN: 201200696W", 9.5, False, B.GREY, 10)
center_text("LEARNER GUIDE", 26, True, B.BLUE, 2)
center_text("For", 11, False, B.GREY, 10)
center_img(os.path.join(B.ASSETS, "cysa_logo.png"), 1.3)
doc.add_paragraph().paragraph_format.space_after = Pt(2)
center_text(C.COURSE_TITLE, 20, True, B.INK, 4)
center_text(f"TGS Ref No: {C.COURSE_CODE}", 11, False, B.INK, 10)
center_text("Conducted by", 10, False, B.GREY, 1)
center_text("Tertiary Infotech Academy Pte Ltd", 12, True, B.INK, 1)
center_text("UEN: 201200696W", 9.5, False, B.GREY, 10)
center_text(f"Version {C.VERSION_DOC}", 12, True, B.BLUE, 4)
doc.add_page_break()

# ---------- Version control ----------
B.heading(doc, "Document Version Control Record", 1)
B.blue_table(doc, ["Version", "Effective Date", "Summary of Changes", "Author"], [
    ["1.0", "6 July 2026",
     "First edition — Learner Guide aligned to the 30 hands-on CySA+ CS0-003 labs (Killercoda + free tools), with step-by-step instructions, reference diagrams and per-lab learning outcomes.",
     "Tertiary Infotech Academy Pte Ltd"],
], widths=[Inches(0.8), Inches(1.2), Inches(3.6), Inches(1.4)])
doc.add_page_break()

# ---------- Course information ----------
B.heading(doc, "Course Information", 1)
B.para(doc, "Welcome! This Learner Guide takes you click-by-click through every hands-on lab in the WSQ "
            f"{C.COURSE_TITLE} (Course Code: {C.COURSE_CODE}). Across four domains you will build practical, "
            "job-ready skills in security operations, vulnerability management, incident response, and reporting — "
            "each mapped to the CompTIA CySA+ CS0-003 exam objectives.")
B.note_box(doc, "Most labs run in the free Killercoda Ubuntu Playground (https://killercoda.com/playgrounds/scenario/ubuntu) — "
                "no local install required. A few use free web tools or free virtual appliances (noted per lab). "
                "Work through the labs in order; each builds on the skills of the one before it.")
B.heading(doc, "Learning Outcomes", 2)
B.para(doc, "By the end of this course, learners will be able to:", bold=True, space=2)
B.bullets(doc, C.LEARNING_OUTCOMES)
B.heading(doc, "Skills Framework Alignment", 2)
B.para(doc, f"TSC Title: {C.TSC_TITLE}   |   TSC Code: {C.TSC_CODE}", bold=True)
B.blue_table(doc, ["Knowledge (K)", "Abilities (A)"],
             [[k, a] for k, a in __import__('itertools').zip_longest(C.TSC_KNOWLEDGE, C.TSC_ABILITIES, fillvalue="")],
             widths=[Inches(3.2), Inches(3.2)])
B.heading(doc, "Course Outline", 2)
for dom in C.DOMAINS:
    B.para(doc, f"{dom['topic']}: {dom['name']} ({dom['weight']}%)", bold=True, color=B.BLUE, space=2)
    B.bullets(doc, dom['concepts'])
B.heading(doc, "Assessment", 2)
B.bullets(doc, ["Written Assessment (SAQ) — 1 hour", "Case Study (CS) — 1 hour",
                "Format: Open Book (Slides, Learner Guide or approved materials only)",
                "Competency: minimum 75% attendance and assessed as 'Competent'."])
B.heading(doc, "How to Use the Hands-On Labs", 2)
B.bullets(doc, [f"Open the Killercoda playground: {C.KILLERCODA}",
                "Follow each step in order; type the commands exactly as shown in the shaded code blocks.",
                "Reset the playground between labs that change firewall rules or install heavy services.",
                "Web-based and local-appliance labs are flagged in each lab's environment note."])
doc.add_page_break()

# ---------- Labs ----------
def render_blocks(blocks):
    for b in blocks:
        if b[0] == 'text': B.para(doc, LP.md_inline(b[1]))
        elif b[0] == 'code': B.code_block(doc, b[1])
        elif b[0] == 'bullets': B.bullets(doc, [LP.md_inline(x) for x in b[1]])
        elif b[0] == 'numbered': B.bullets(doc, [LP.md_inline(x) for x in b[1]], style='List Number')
        elif b[0] == 'subhead': B.para(doc, LP.md_inline(b[1]), bold=True, color=B.GREY, space=3)

for dom in C.DOMAINS:
    B.heading(doc, f"{dom['topic']}: {dom['name']}", 1)
    B.para(doc, f"CompTIA CySA+ CS0-003 exam weighting: {dom['weight']}%. Hands-on Labs {dom['labs'][0]}–{dom['labs'][-1]}.",
           italic=True, color=B.GREY)
    for ln in dom['labs']:
        L = LABS[ln]
        B.heading(doc, LP.md_inline(L['title']), 2)
        for b in L['intro']:
            if b[0] == 'text': B.para(doc, LP.md_inline(b[1]))
        B.note_box(doc, C.ENV[ln], label="Environment:")
        if ln in C.DIAGRAMS:
            for f in C.DIAGRAMS[ln]:
                src = os.path.join(IMG, f)
                if os.path.exists(src):
                    B.image_centered(doc, src, 5.4)
        for sec in L['sections']:
            hd = "What You Learned" if sec['heading'].lower().startswith('what you learned') else LP.md_inline(sec['heading'])
            B.heading(doc, hd, 3)
            render_blocks(sec['blocks'])
        doc.add_page_break()

B.add_footer(doc)
lg_name = f"LG-{C.COURSE_TITLE}.docx"
doc.save(os.path.join(OUT_DIR, lg_name))
print("Saved LG DOCX:", lg_name)
